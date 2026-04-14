import docx
from docx import Document
import os

def read_word_document(file_path):
    """读取Word文档内容"""
    try:
        doc = Document(file_path)
        full_text = []

        for para in doc.paragraphs:
            full_text.append(para.text)

        # 读取表格内容
        tables_text = []
        for i, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                table_text.append(" | ".join(row_text))
            tables_text.append(f"表格 {i+1}:\n" + "\n".join(table_text))

        return "\n".join(full_text), tables_text
    except Exception as e:
        return f"Error reading document: {str(e)}"

# 文件路径
file_path = "D:\\AI工作区\\conversations\\conv-1775790990474\\1775791026848-创新发展专项基金申报书 大学生专项目-207010108仵煜 最新版.doc"

# 读取文档
result = read_word_document(file_path)

if isinstance(result, tuple):
    text_content, tables = result
    print("=== 文档文本内容 ===")
    print(text_content)
    print("\n" + "="*50 + "\n")

    print("=== 表格内容 ===")
    for table in tables:
        print(table)
        print("\n" + "-"*30 + "\n")
else:
    print(result)