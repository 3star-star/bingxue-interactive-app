import docx
import os

def extract_docx_text(file_path):
    """提取Word文档中的所有文本"""
    try:
        doc = docx.Document(file_path)
        full_text = []

        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                full_text.append(" | ".join(row_text))

        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading file: {str(e)}"

# 使用示例
file_path = r"D:\AI工作区\conversations\conv-1775791118793\1775791162014-创新发展专项基金申报书 大学生专项目-207010108仵煜 最新版.docx"
text = extract_docx_text(file_path)
print(text)